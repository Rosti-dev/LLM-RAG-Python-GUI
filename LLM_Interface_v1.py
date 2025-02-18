#Version der GUI v1.0

import os
import tkinter as window
from dotenv import load_dotenv
import datetime
from tkinter import filedialog, messagebox, scrolledtext
import json
import requests
import threading
from docx import Document
import re

#Anforderungen um das Skript inklusive Abhängigkeiten zu starten sind der requirements.txt Datei zu entnehmen.
#Für den Betrieb des Vektordatenservers ist die Installation von der "Visualstudio cpp build tools app" nötig
#Für Visualstudio Programm Installation siehe https://visualstudio.microsoft.com/visual-cpp-build-tools/

#Dokumentiert die Antworten der Vektordatenbank, LLM-Server, LLM Anfragen in der Konsole und in Text Datei.
Debug = True

System_Role = "Du bist ein hilfreicher Assistent. Antworte ausschließlich auf die letzte neuste Anfrage des 'user'. Antworte immer ohne Formatierung des Textes."
#Old SystemRole "Du bist ein hilfreicher Assistent. Antworte ausschließlich auf user Aufgabe. Der Background-Context und Previous-User-Input stellen zusätzliche Informationen zu Verfügung."} #You are a helpful assistant.

#Lädt Konfiguration inklusive Parametern und API-Schlüssel.
load_dotenv()


#Lokale LLM Modelle
#.env Modelle für lokale LLM entnehmen. Kleines Default Modell als Standardwert falls Variable nicht gefunden.
LOCAL_LLM_MODELS = os.getenv("LLM_MODELS", "llama3.2:3b,deepseek-r1:1.5b")
LOCAL_LLM_MODELS = LOCAL_LLM_MODELS.split(",")


#Vektordatenbank Parameter - Bitte bei Verwendung eines anderen Vektordatenbankservers anpassen.
VECTORDATABASE_URL = os.getenv("VECTORDB_URL", "http://127.0.0.1:8000") #URL
VECTORDATABASE_UPLOAD = f"{VECTORDATABASE_URL}/upload"              #URL Adresse für das Hochladen von Anhängen.
VECTORDATABASE_REQUEST = f"{VECTORDATABASE_URL}/ask"                #URL Adresse für Anfragen

#Parameter und Default Parameter, falls nicht in der .env festgelegt.
#Bei erstmaliger Ausführung wird eine .env Datei erstellt. Dort können die Parameter festgelegt werden.
LOCAL_IP = os.getenv("LOCAL_IP", "127.0.0.1")
PORT = os.getenv("PORT", "11434")
LLM_SERVER_URL = f"http://{LOCAL_IP}:{PORT}/api/chat"

#OpenAI Parameter - Modelle und API Schlüssel aus .env entnehmen.
OPENAI_API_KEY = os.getenv('OPENAI_KEY', '').strip(" ")
#Modelle (ohne API Abfrage der aktuell verfügbaren Modelle):
OPENAI_MODELS = os.getenv('OPENAI_MODELS', 'gpt-4o-mini,gpt-4o,gpt-4,gpt-4-turbo,o1-mini,o1,o1-pro')
OPENAI_MODELS = OPENAI_MODELS.split(",")

#########

#Protokolliert die Vektordatenbank Ergebnisse, falls "Debug" als wahr gesetzt wurde.
def log_vector_response(vectorserver_response):
    os.makedirs("Vektordatenbank", exist_ok=True)
    log_path = os.path.join("Vektordatenbank", "AnfragenLog.txt")
    formatted_current_time = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(log_path, "a", encoding="utf-8") as f:
        f.write(f"\n\n#### {formatted_current_time} ####\n") #Fomartierung - Unterscheidung zwischen den Ergebnissen.
        f.write(json.dumps(vectorserver_response)) #Json Antwort in txt dokumentieren.  #, ensure_ascii=False, indent=2

########
#Funktionen zur Kommunikation mit LLM und Vektordatenbank

#Anfrage Funktion.
#Zweiter Parameter gibt die Search kwargs an - Ebenfalls für ein besseres Verständnis über den Parameter als "max_results" gesetzt.
#Siehe Flask Server Code für mehr Kontext.
def request_vector_context(question, max_results=5):
    try:
        #Anfragenformatierung und Konfiguration über die gewünschte Menge von Ergebnissen.
        payload = {"question": question, "max_results": max_results }
        response = requests.post(f"{VECTORDATABASE_URL}/ask", json=payload) #URL und Inhalt anpassen: aktuelle Konfiguration is mit VektordatenbankServer.py zu nutzen.
        response.raise_for_status()
        vector_response = response.json()
        return vector_response
    except requests.exceptions.RequestException as e:
        return {"error": f"Fehler bei der Kommunikation mit der Vektordatenbank: {str(e)}"}

#Vorbereitung des Antwort-Ergebnistextes - Prüfung auf doppelte Ergebnisse
def prepare_context(documents):
    unique_contents = set()
    structured_context = []
    count = 0
    for doc in documents:
        content = doc.get("content", "").strip()
        if content != "" and content not in unique_contents:
            unique_contents.add(content)
            count += 1
            structured_context.append(f"{count}. {content}")
    return "\n".join(structured_context)

def clean_local_llm_response(response): #Reasoning bei Modellen wie Deepseek-R1 nicht in der Antwort darstellen.
    #Entfernt alle Inhalte zwischen <think>...</think> in der Antwort.
    cleaned_response = re.sub(r"<think>.*?</think>", "", response, flags=re.DOTALL)
    return cleaned_response.strip()

#Funktion zum Adressieren einer lokalen/selbstgehosteten Ollama Instanz
#Paramter Nutzung ergibt sich aus deren Bezeichnung:
#Conversation_history: Gesprächsverslauf mit Modell; Question: Aufgabe/Fragestellung an LLM, model: LLM Name, Vector: Anfragenbezogener Kontext für RAG Implementation, leer falls nicht vorhanden
def ask_llm_server(conversation_history, question, model, vector_context_result=""):
    #Unterschiedliche JSON payload, falls keine Kontext Information vorhanden, dann wird der Teil der Anfrage nicht inkludiert.
    if vector_context_result != "":
        background_context = {"role": "Context", "content": f"Relevanter Kontext:\n{vector_context_result}"}
        payload = {"model": model,
            "messages": conversation_history + [background_context] + [{"role": "user", "content": question}]}
    else:
        payload = {"model": model,
            "messages": conversation_history + [{"role": "user", "content": question}]}
    try:
        #Fix: Ohne Ensure ascii sind die Umlaute unlesbar
        if Debug:
            print("JSON Anfrage an den lokalen LLM-Server:\n\n" + json.dumps(payload, ensure_ascii=False, indent=2) + "\n\n")
        response = requests.post(LLM_SERVER_URL, json=payload, stream=True)
        """Streaming ist zwar nicht auf der GUI unterstützt, hilft dennoch da beim Schließen der GUI bei 
        existierender Anfrage an den LLM Server auch die Verarbeitung der abgebrochen wird. Streaming wäre wünschenswert, 
        um Fortschritt langsamer LLM Ausgaben beobachten zu können. Modelle wie Deepseek und o3-mini haben eine erhöhte Verarbeitungszeit."""

        complete_message = ""
        #JSON parsen
        for line in response.iter_lines():
            if line: #Leere Zeilen filtern.
                chunk = json.loads(line)
                complete_message += chunk.get("message", {}).get("content", "")
        complete_message = clean_local_llm_response(complete_message) #clean_local_llm_resonse: Reasoning wird an dieser Stelle entfernt (z.B. wie bei deepseek-r1:1.5b Modellen)
        return {"response": complete_message}
    #Fehlermeldung bei Verbindungsverlust, zu großer Anfrage etc.
    except requests.exceptions.RequestException as e:
        return {"error": f"Fehler bei der Kommunikation mit dem LLM-Server: {e}"}

#In zukünftiger Version Parametrisieren, um es für weitere Plattformen nutzen zu können oder duplizieren und an eine Plattform anpassen.
def ask_openai_server(conversation_history, documents, question, model="gpt-4o-mini"):
    """Funktion zur Kommunikation mit OPENAI servern (ChatGPT) - Mit geänderter URL kann man damit alle OpenAI ähnlichen API Interfaces,
    wie zum Beispiel Openrouter verwenden, welches eine Vielzahl von LLMs betreibt."""
    #API Key vorhanden?
    if not OPENAI_API_KEY:
        return {"error": "Kein gültiger OpenAI API Schlüssel gefunden. Bitte OPENAI_KEY in der .env angeben."}
    #Bei Fehlermeldungen des Vektorservers wird keine Anfrage an OpenAI gestellt. Optional entfernen.
    if "error" in documents:
        print("Error in der Antwort des Vektorservers gefunden.")
        return documents

    #Kontext des Vektorservers
    if documents["documents"]:
        context = prepare_context(documents["documents"])
    else:
        context = None

    #Konversationsverlauf duplizieren, um conversation_history nicht zu verändern und den Verlauf für die neue Anfrage zu nutzen.
    messages = conversation_history.copy()

    #Kontext als System Nachricht angehangen. Prüfe ideale Umsetzung, da es mehere Möglichkeiten gibt: kann ebenfalls in der Anfrage angehangen werden.
    if context:
        messages.append({"role": "system", "content": f"Zusätzlicher Kontext: {context}"})
    #Neue Anfrage anhängen.
    messages.append({"role": "user", "content": question})
    #Header zur Angabe des Schlüssels.
    headers = {
        "Content-Type": "application/json", #Inhalt, kann alternativ auch Audio/Bild oder andere Formate beinhalten.
        "Authorization": f"Bearer {OPENAI_API_KEY}"
    }

    data = {"model": model, "messages": messages } #Modellauswahl und Anfrageninhalt zusammenfassen.

    if Debug:
        print("JSON Anfrage an den OpenAI Server:\n\n" + json.dumps(data, ensure_ascii=False, indent=2) + "\n\n")

    try:
        response = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=200
        )
        response.raise_for_status() #Fehler in der Antwort erkennen.
        result = response.json()
        if Debug:
            print("JSON Antwort des OpenAI Servers:\n\n" + json.dumps(result, ensure_ascii=False, indent=2) + "\n\n")

        return {"response": result["choices"][0]["message"]["content"]} #Ergebnis zurückgeben.
    except requests.exceptions.RequestException as e:
        return {"error": f"Fehler bei der Kommunikation mit OpenAI: {e}"}


#Uploadfunktion zur Vektordatenbank
#Zukünftig wäre es besser über https zu posten und einen API Schlüssel zu erstellen, um Zugriff auf sensible Daten zu verhindern.
def upload_files_to_rag_server(file_paths):
    """Im folgenden werden die Funktion zum Hochladen von Text, Docx und PDF Dateien definiert und konfiguriert."""
    if not file_paths:
        return "Keine Dateien ausgewählt."
    try:
        for file_path in file_paths:
            with open(file_path, 'rb') as f: #Datei als Binary öffnen
                response = requests.post(VECTORDATABASE_UPLOAD, files={"file": f})
            response.raise_for_status()
        return "Hochladen erfolgreich"
    except requests.exceptions.RequestException as e:
        return f"Fehler beim Hochladen der Datei: {e}"

#####
#Im Folgenden werden zum größten Teil die Komponenten für die Benutzeroberfläche erstellt.
#####

#Frame für Box der Modelle das scrollbar ist.
def create_scrollable_frame(parent):
    #Canvas und Frame/Rahmen erstellen
    canvas = window.Canvas(parent, borderwidth=0, height=150)
    frame = window.Frame(canvas)
    #Scroll Leiste
    vsb = window.Scrollbar(parent, orient="vertical", command=canvas.yview)
    canvas.configure(yscrollcommand=vsb.set)
    vsb.pack(side="right", fill="y")
    canvas.pack(side="left", fill="both", expand=True)
    canvas.create_window((0, 0), window=frame, anchor="nw")
    #Funktion zur Aktualisierung des Canvas.
    def onFrameConfigure(event):
        canvas.configure(scrollregion=canvas.bbox("all"))
    frame.bind("<Configure>", onFrameConfigure)
    return frame


#Klasse zur Visualisierung der Nutzeroberfläche
class LLMApp(window.Tk):
    def __init__(self):
        super().__init__()
        self.title("LLM Chatbot Interaktionsoberfläche")
        self.geometry("1200x720")  #Fenstergröße nach Bedarf anpassen oder zukünftig parametrisieren.

        #Liste für die Dokumente die hochgeladen werden.
        self.attached_files = []

        self.llm_conversation_history = {}
        all_models = LOCAL_LLM_MODELS + OPENAI_MODELS #Listen der LLM Modelle kombinieren.
        for model_name in all_models:
            system_message = {
                "role": "system",
                "content": System_Role
            }

        #Erstellt individuelle Historie für jede Modellbezeichnung mit Systemnachricht am Anfang des Verlaufs.
        self.llm_conversation_history[model_name] = [system_message]



        ###########################################
        #Layout
        ##########################################
        #Erstellt PanedWindow - geteiltes Fenster, dass sich horizontal verschieben lässt.
        self.paned = window.PanedWindow(self, orient=window.HORIZONTAL)
        self.paned.pack(fill=window.BOTH, expand=True)

        #Linker Bereich
        self.left_frame = window.Frame(self.paned, width=200) #Erstellt Frame für Oberflächenelemente
        self.paned.add(self.left_frame)                       #Fügt den linken Frame hinzu

        #Rechter Bereich
        self.right_frame = window.Frame(self.paned)           #Erstellt Frame für Oberflächenelemente
        self.paned.add(self.right_frame)                      #Fügt den rechten Frame hinzu
        self.paned.paneconfigure(self.right_frame, minsize=500)

        #Teilt die linke Seite vertikal ein.
        self.left_paned = window.PanedWindow(self.left_frame, orient=window.VERTICAL)
        self.left_paned.pack(fill=window.BOTH, expand=True)

        #Teil oben links. Eingabefeld und Schaltflächen.
        self.left_top = window.Frame(self.left_paned)
        self.left_paned.add(self.left_top, stretch="always")  #"always" - dehnt sich mit dem Fenster

        #Teil links unten. Modellauswahl
        self.left_bottom = window.Frame(self.left_paned)
        self.left_paned.add(self.left_bottom)


        ###########
        #Eingabefeld (links oben)
        window.Label(self.left_top, text="Eingabefeld:").grid(row=0, column=0, sticky="nw") #Beschriftung
        #Eingabefeld mit Scrollleiste
        self.input_text = scrolledtext.ScrolledText(self.left_top, height=4)
        self.input_text.grid(row=1, column=0, columnspan=4, pady=5, sticky="nsew") #Ausrichtung, Position
        self.left_top.rowconfigure(1, weight=1)  #Zeile dehnt sich aus.
        self.left_top.columnconfigure(0, weight=1) #Spalte dehnt

        #Schaltfläche zum Senden von Anfragen.
        self.send_button = window.Button(self.left_top, text="Senden", command=self.send_request)
        self.send_button.grid(row=4, column=3, sticky="e", padx=5, pady=5)

        ##################
        #Dateiupload Bereich
        ##
        #Rahmen
        upload_frame = window.LabelFrame(self.left_top, text="Datei-Upload Vektorserver")
        upload_frame.grid(row=3, column=0, columnspan=4, padx=5, pady=2, sticky="ew")
        #Liste ausgewählter Dateien inklusive Dateipfad
        self.files_list = window.Listbox(upload_frame, height=3)
        self.files_list.pack(side=window.LEFT, fill=window.BOTH, expand=True, padx=2, pady=2)
        #Schalftlächen Container
        file_buttons_frame = window.Frame(upload_frame)
        file_buttons_frame.pack(side=window.LEFT, fill=window.Y, padx=2, pady=2)
        #Dateiauswahl Schalftläche
        self.attach_button = window.Button(file_buttons_frame, text="Dateien auswählen", command=self.select_files)
        self.attach_button.pack(anchor="nw", pady=2)
        #Schaltfläche Datei hochladen durchführen
        self.upload_button = window.Button(file_buttons_frame, text="Dateien hochladen", command=self.upload_files)
        self.upload_button.pack(anchor="nw", pady=2)
        #Schaltfläche leert Dateiliste
        self.clear_button = window.Button(file_buttons_frame, text="Dateien deselektieren", command=self.clear_file_selection)
        self.clear_button.pack(anchor="nw", pady=2)

        #Kontrollkästchen<->Checkboxen - Anfrage behalten verhindert das Löschen
        self.retain_input_var = window.BooleanVar(value=True)
        self.retain_input_checkbox = window.Checkbutton(self.left_top, text="Anfrage behalten", variable=self.retain_input_var)
        self.retain_input_checkbox.grid(row=4, column=0, sticky="w", padx=5, pady=2)

        #Kontrollkästchen "Konversation automatisch speichern"
        self.auto_save_var = window.BooleanVar(value=False)
        self.auto_save_checkbox = window.Checkbutton(self.left_top, text="Konversation automatisch speichern", variable=self.auto_save_var)
        self.auto_save_checkbox.grid(row=5, column=0, sticky="w", padx=5, pady=2)

        #Kontrollkästchen Vektordatenbank einbeziehen.
        self.use_vektordb_variables = window.BooleanVar(value=False)
        self.use_vektordb_checkbox = window.Checkbutton(self.left_top, text="Vektordatenbank einbeziehen", variable=self.use_vektordb_variables)
        self.use_vektordb_checkbox.grid(row=6, column=0, sticky="w", padx=5, pady=2)

        #Verbindung mit lokalem LLM Server und Vektorserver testen. -> Ausgabe der Resultate via Messagebox.
        self.test_button = window.Button(self.left_top, text="Verbindung testen", command=self.run_test_connection)
        self.test_button.grid(row=7, column=0, sticky="w", padx=5, pady=5)

        #Modelle Auswahlbereiche die scrollbar sind. (Noch nicht mit Scrollmauszeiger bedienbar aber vorerst fallen bei einer langen Liste außerhalb des Fensters.
        models_frame = window.Frame(self.left_bottom)
        models_frame.pack(fill=window.BOTH, expand=True, padx=5, pady=5)
        #Lokale LLM-Modelle auf dem Grid eingesetzt.
        local_models_labelframe = window.LabelFrame(models_frame, text="Lokale Modelle", width=75)
        local_models_labelframe.grid(row=0, column=0, sticky="nsew", padx=5)
        models_frame.grid_columnconfigure(0, weight=1)

        #Lokale Modelle Scrollbar
        local_models_scrollable = create_scrollable_frame(local_models_labelframe)
        self.local_model_vars = []
        for model in LOCAL_LLM_MODELS:
            choice_state = window.BooleanVar() #Auswahlzustand
            cb = window.Checkbutton(local_models_scrollable, text=model, variable=choice_state, command=self.check_model_limits)
            cb.pack(anchor="w")
            self.local_model_vars.append((model, choice_state))

        #OpenAI Modelle Auswahl- Werden nur angezeigt wenn API Key vorhanden ist.
        #ToDo: Hinzufügen einer automatischen Modellabfrage zu Beginn des Codes über den API Key, damit
        #eine Auswahl aktueller Modelle vorliegt.
        if OPENAI_API_KEY:
            openai_models_labelframe = window.LabelFrame(models_frame, text="OpenAI Modelle", width=75)
            openai_models_labelframe.grid(row=0, column=1, sticky="nsew", padx=5)
            models_frame.grid_columnconfigure(1, weight=1)
            openai_models_scrollable = create_scrollable_frame(openai_models_labelframe)
            self.openai_model_vars = []
            for model in OPENAI_MODELS:
                choice_state = window.BooleanVar()
                cb = window.Checkbutton(openai_models_scrollable, text=model, variable=choice_state, command=self.check_model_limits)
                cb.pack(anchor="w")
                self.openai_model_vars.append((model, choice_state))

        #Rechte Hälfte Ausgabefelder.
        #Erstes Ausgabefeld mit Beschriftung.
        self.label_output1 = window.Label(self.right_frame, text="Ausgabe (kein Modell)")
        self.label_output1.pack()
        self.output_text_1 = scrolledtext.ScrolledText(self.right_frame, height=8)
        self.output_text_1.pack(fill=window.BOTH, expand=True)
        #Zweites Ausgabefeld
        self.label_output2 = window.Label(self.right_frame, text="Ausgabe (kein Modell)")
        self.label_output2.pack()
        self.output_text_2 = scrolledtext.ScrolledText(self.right_frame, height=8)
        self.output_text_2.pack(fill=window.BOTH, expand=True)
        #Drittes Ausgabefeld
        self.label_output3 = window.Label(self.right_frame, text="Ausgabe (kein Modell)")
        self.label_output3.pack()
        self.output_text_3 = scrolledtext.ScrolledText(self.right_frame, height=8)
        self.output_text_3.pack(fill=window.BOTH, expand=True)

        ########################
        #Schaltflächen unten rechts: Konversationsverlauf speichern & Konversationsverlauf öffnen (Word Datei)
        ###########
        #Schalftllächen zur
        btn_frame = window.Frame(self.right_frame)
        btn_frame.pack(fill=window.X, pady=5)

        # Bildet im in der aktuellen Version keine sinnvolle Funktionalität ab.
        """
        #Erstes Ausgabefeld wird in die Zwischenablage kopiert per Schalftlächenbestätigung.
        #self.copy_button = window.Button(btn_frame, text="Erstes Ausgabefeld kopieren", command=self.copy_last_message)
        #self.copy_button.pack(side=window.LEFT, padx=5)
        """
        #Konversation speichern Schaltfläche
        self.save_button = window.Button(btn_frame, text="Konversation speichern", command=self.save_conversation_to_word)
        self.save_button.pack(side=window.LEFT, padx=5)
        #Konversationsverlauf öffnen
        self.open_button = window.Button(btn_frame, text="Konversationsverlauf öffnen", command=self.open_conversation)
        self.open_button.pack(side=window.LEFT, padx=5)

        #Konversationshistorie - Liste zum Speichern des Chatverlaufs jedes einzelnen Modells
        self.conversation_history = []

    ################
    #Methoden
    ############
    #Dateiupload Vektorserver - selektion der Dateien.
    def select_files(self):
        #Dateiauswahl & speichern der vollständigen Pfade.
        #Vorerst werden pdf,txt,docx unterstützt.
        paths = filedialog.askopenfilenames(filetypes=[("All Files", "*.pdf;*.docx;*.txt")])
        for path in paths:
            absolute_path = os.path.abspath(path)
            self.attached_files.append(absolute_path)
            self.files_list.insert(window.END, absolute_path)
    #Dateiupload  durchführen
    def upload_files(self):
        if self.attached_files:
            threading.Thread(target=self.upload_sequentially, daemon=True).start()
        else:
            messagebox.showinfo("Upload", "Keine Dateien zum Hochladen ausgewählt.")
            return
    #Dateien nacheinander hochladen (Fix)
    def upload_sequentially(self):
        files_to_upload = list(self.attached_files)
        for fpath in files_to_upload:
            result = upload_files_to_rag_server([fpath])
            if "erfolgreich" in result.lower():
                self.remove_file_from_list(fpath)
            else:
                messagebox.showerror("Upload-Fehler", result)
        messagebox.showinfo("Upload", "Upload-Prozess abgeschlossen.")
    #Funktion - Dateien von Liste entfernen
    def remove_file_from_list(self, filepath):
        if filepath in self.attached_files:
            self.attached_files.remove(filepath)
        for i in range(self.files_list.size()):
            if self.files_list.get(i) == filepath:
                self.files_list.delete(i)
                break
    #2. Funktion Dateien von Liste entfernen
    def clear_file_selection(self):
        self.attached_files.clear()
        self.files_list.delete(0, window.END)


    ###########
    #Verbindungstest - Konfiguration und Addressen verifizieren
    def run_test_connection(self):
        #Verbindungstest im Thread starten
        threading.Thread(target=self.test_connection, daemon=True).start()

    def test_connection(self):
        #Verbindungstest zu lokalem LLM Server und Vektordatenbank
        status_message = ""
        #Test LLM-Server mit leerer Anfrage, warte auf Antwort 404
        try:
            resp = requests.get(LLM_SERVER_URL, timeout=5)
            if resp.status_code == 404:
                status_message += "Erfolgreich lokaler LLM-Server. "
            else:
                status_message += "Fehlgeschlagen lokaler LLM-Server. "
        except requests.RequestException:
            status_message += "Fehlgeschlagen LLM Server. "

        #Test: Vektordatenbank
        try:
            resp2 = requests.get(VECTORDATABASE_URL, timeout=5)
            if resp2.status_code == 200:
                status_message += "Erfolgreich Vektordatenbank. "
            else:
                status_message += "Fehlgeschlagen Vektordatenbank. "
        except:
            status_message += "Fehlgeschlagen Vektordatenbank. "
        messagebox.showinfo("Verbindung Testergebnis", status_message)

    ####################
    #Logik Modelle
    ####################
    #Ausgewählte Modellanzahl auf 3 begrenzen (Anzahl der Ausgabebereiche.
    def check_model_limits(self):
        total_selected = 0
        #Zähle die Anzahl der Modelle, die Augewählt wurden
        #In der Iteration wird der Mdellname durch "_" ignoriert.
        for _, model_selection_state in self.local_model_vars:
            if model_selection_state.get():
                total_selected += 1
        for _, model_selection_state in self.openai_model_vars:
            if model_selection_state.get():
                total_selected += 1
        if total_selected > 3:
            messagebox.showerror("Fehler", "Es können maximal 3 Modelle gleichzeitig ausgewählt werden.")
            #Zuletzt getätigte Modellauswahl zurücksetzen
            for model, model_selection_state in reversed(self.openai_model_vars + self.local_model_vars):
                if model_selection_state.get():
                    model_selection_state.set(False)
                    break

    ####################################
    #Anfragenlogik
    ####################################
    def send_request(self):
        #Liest Nutzereingabe
        user_input = self.input_text.get("1.0", window.END).strip()
        if user_input == "":
            messagebox.showwarning("Fehler", "Das Eingabefeld ist leer.")
            return

        #Liste der lokalen Modelle
        selected_local_models = []
        #Iteriert durch die Modelle und deren Auswahlzustand
        for model_name, model_state in self.local_model_vars:
            #Fügt Namen bei ausgewähltem Modell an Ausgabe an.
            if model_state.get():
                selected_local_models.append(model_name)

        #Liste der Openai Modelle
        selected_openai_models = []
        #Iteriert durch die Modelle und deren Auswahlzustand
        for model_name, model_state in self.openai_model_vars:
            if model_state.get():
                selected_openai_models.append(model_name)

        #Kombiniert Listen
        all_models = selected_local_models + selected_openai_models

        #Setzt die Namen der Modelle zu den entsprechenden Ausgabefeldern.
        self.setup_output_labels(all_models)

        #Anfragen an Modelle im Thread senden.
        threading.Thread(target=self.process_request_sequential, args=(user_input, all_models), daemon=True).start()

    #Aktualisiert die Namen der Ausgabefelder zu den entsprechenden Modellnamen
    def setup_output_labels(self, model_list):
        #Ausgabe 1
        if len(model_list) > 0:
            self.label_output1.config(text=f"Ausgabe (1. Modell = {model_list[0]})")
        else:
            self.label_output1.config(text="Ausgabe (kein Modell)")
        self.output_text_1.delete("1.0", window.END)

        #Ausgabe 2
        if len(model_list) > 1:
            self.label_output2.config(text=f"Ausgabe (2. Modell = {model_list[1]})")
        else:
            self.label_output2.config(text="Ausgabe (kein Modell)")
        self.output_text_2.delete("1.0", window.END)

        #Ausgabe 3
        if len(model_list) > 2:
            self.label_output3.config(text=f"Ausgabe (3. Modell = {model_list[2]})")
        else:
            self.label_output3.config(text="Ausgabe (kein Modell)")
        self.output_text_3.delete("1.0", window.END)

    #Sequentielle Abfrage der ausgewählten Modelle.
    #Optimierungsvorschlag: OpenAI Modelle und andere schnelle API Anbieter, die eine schnellere Verarbeitungszeit aufweisen sollten hier priorisiert werden,
    #   da dies dem Nutzer die Möglichkeit bietet bereits die Inhalte der Ausgabe zu evaluieren.
    def process_request_sequential(self, user_input, model_list):
        """ Anfrage an die Vektordatenbank und die entsprechenden LLM Server.
        Abschließend werden in der Funktion die Ergebnisse der LLM Anfragen in den Ausgabefeldern eingesetzt.  """
        #Anfrage Vektordatenbank, falls aktiviert.
        vectordb_result = {"documents": []}
        vector_context_result = ""
        if self.use_vektordb_variables.get():
            vectordb_result = request_vector_context(user_input, max_results=5)
            if Debug:
                log_vector_response(vectordb_result)
            #Kontext für die LLM Modelle, falls aktiviert und kein Fehler.
            if "documents" in vectordb_result and not "error" in vectordb_result:
                vector_context_result = prepare_context(vectordb_result["documents"])

        responses = [] #Liste LLM Antworten
        for idx, model_name in enumerate(model_list):
            if model_name in OPENAI_MODELS:
                #Lädt den bisherigen Gesprächsverlauf, falls vorhanden, ansonsten initialisiert Gesprächsverlauf
                history = self.llm_conversation_history.get(model_name, [{"role": "system", "content": System_Role}])
                resp = ask_openai_server(history, vectordb_result, user_input, model=model_name)
                #Fügt Rollen und Verlauf in die Historie hinzu.
                history.append({"role": "user", "content": user_input})
                history.append({"role": "assistant", "content": resp.get("response", "")})
                self.llm_conversation_history[model_name] = history

            elif model_name in LOCAL_LLM_MODELS:
                #Lokale Modelle mit Gedächtnis
                history = self.llm_conversation_history.get(model_name, [{"role": "system", "content": System_Role}])
                resp = ask_llm_server(conversation_history=history, question=user_input, model=model_name, vector_context_result=vector_context_result)
                history.append({"role": "user", "content": user_input})
                history.append({"role": "assistant", "content": resp.get("response", "")})
                self.llm_conversation_history[model_name] = history

            responses.append(resp["response"])
            #Ergebnis anzeigen.
            self.update_output_box(idx, responses[-1])

        def update_output_box(self, model_index, text):
            # Fügt die LLM Ergebnisse in die entsprechenden Augabefelder ein.
            if model_index == 0:
                self.output_text_1.delete("1.0", window.END)
                self.output_text_1.insert(window.END, text)
            elif model_index == 1:
                self.output_text_2.delete("1.0", window.END)
                self.output_text_2.insert(window.END, text)
            elif model_index == 2:
                self.output_text_3.delete("1.0", window.END)
                self.output_text_3.insert(window.END, text)

        #Konversationsergebnisse für Logging sammeln, strukturieren und Datum hinzufügen.
        self.conversation_history.append({
            "input": user_input,
            "models": model_list,
            "responses": responses,
            "time": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        })

        #Automatisch Konversation speichern, falls aktiviert.
        if self.auto_save_var.get():
            self.save_conversation_to_word()

        #Eingabefeld geleert, falls Checkbox nicht aktiviert.
        if not self.retain_input_var.get():
            self.input_text.delete("1.0", window.END)



    ##################################################################
    #Weitere Hilfsfunktionen
    #Speichert Konversation in Docx/Word Datei.
    def save_conversation_to_word(self):
        doc_path = "Konversationsverlauf_LLM.docx"
        if os.path.exists(doc_path):
            document = Document(doc_path)
        else:
            document = Document()
            document.add_heading('Konversationsverlauf LLM', 0)

        now = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        #Letzter Eintrag der Konversationshistorie
        if self.conversation_history:
            latest_entry = self.conversation_history[-1]
            document.add_paragraph(f"Datum und Zeit: {now}")
            document.add_paragraph(f"Gewählte Modelle: {', '.join(latest_entry['models'])}")
            document.add_paragraph(f"Eingabe:\n{latest_entry['input']}")
            for i, resp in enumerate(latest_entry['responses']):
                document.add_paragraph(f"Antwort Modell {i+1}:\n{resp}")
            document.add_page_break()
            document.save(doc_path)
            if Debug:
                print(f"Konversation wurde in '{doc_path}' gespeichert.")
        else:
            messagebox.showinfo("Info", "Keine Konversation zum Speichern vorhanden.")
    #Öffnet die Docx/Word Datei, die den Gesprächsverlauf beinhaltet
    def open_conversation(self):
        doc_path = "Konversationsverlauf_LLM.docx"
        if os.path.exists(doc_path):
            os.startfile(doc_path)
        else:
            messagebox.showwarning("Datei nicht gefunden", "Es gibt noch keine gespeicherte Konversation.")

#Start der App
if __name__ == "__main__":
    app = LLMApp()
    app.mainloop()
